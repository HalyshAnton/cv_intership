import re
import warnings
import argparse

from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import cv2
import pymupdf
import layoutparser as lp
import numpy as np
import pytesseract
import easyocr
import torch
from docx import Document
from tqdm import tqdm

from data_preprocessor import preprocess_image
from read_pdf import read_pdf
from model import TextClassifier

import collections
collections.Iterable = collections.abc.Iterable

pytesseract.pytesseract.tesseract_cmd = r'C:\Programs\cv_intership\Tesseract-OCR\tesseract.exe'

warnings.filterwarnings('ignore')

text_classifier = TextClassifier()
text_classifier.load_state_dict(torch.load('weights.pt',
                                           map_location=torch.device('cpu')))
text_classifier.eval()

reader = easyocr.Reader(['en'])

processor = TrOCRProcessor.from_pretrained("microsoft/trocr-base-handwritten")
trocr_model = VisionEncoderDecoderModel.from_pretrained("microsoft/trocr-base-handwritten")

model = lp.Detectron2LayoutModel('lp://PubLayNet/mask_rcnn_X_101_32x8d_FPN_3x/config',
                                 extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.4],
                                 label_map={0: "Text", 1: "Title", 2: "List",
                                            3: "Table", 4: "Figure"})

def merge_blocks(layout, thresh_width=5):
    """
    Merge intersecting lp.TextBlock from layout

    Args:
        layout (lp.Layout): layout with y-coordinate
        sorted TextBlocks

        thresh_width (int): minimum intersecting width
        in pixels needed to merge 2 TextBlock

    Return:
        lp.Layout: layout with merged blocks
    """
    if len(layout) == 0:
        return layout

    blocks = []
    current_block = layout[0]

    for block in layout:
        inter = current_block.intersect(block)
        inter_width = inter.block.y_2 - inter.block.y_1

        if inter_width > thresh_width:
            current_block = current_block.union(block)
        else:
            blocks.append(current_block)
            current_block = block

    blocks.append(current_block)
    return lp.Layout(blocks)


def get_text(block, img):
    """
    Get text from image using pytesseract

    Args:
        block (lp.TextBlock): image block returned by lp model

        img (np.ndarray): whole image from what the block will be cut

    Returns:
        str: text from image block
    """

    segment_image = (block
                     .pad(left=5, right=5, top=5, bottom=5)
                     .crop_image(img))

    is_handwritten = text_classifier.predict(segment_image)

    if is_handwritten:
        text = get_handwritten_text(segment_image)
    else:
        text = pytesseract.image_to_string(segment_image, config='--psm 6')
        text = text.replace('\n', ' ')

    return text, is_handwritten


def get_handwritten_text(image):
    """
    Extract handwritten text from an image using image processing and OCR.

    Args:
        image (np.ndarray): The input image from which handwritten text is
        to be extracted. The image should be in RGB format.

    Returns:
        str: The extracted handwritten text from the image
    """

    img = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
    img = cv2.medianBlur(img, 5)
    img = cv2.adaptiveThreshold(
        img, 255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        11, 5)

    hist = img.mean(axis=1)

    cuts = []  # indexes of lines edges
    i = 0

    while i < len(hist):
        if hist[i] > 254:
            begin = i
            while i < len(hist) and hist[i] > 254:
                i += 1
            cuts.append((i + begin) // 2)  # middle of the whitespace region
        else:
            i += 1

    text = ''

    # trocr_model works with one line text
    for lower, upper in zip(cuts, cuts[1:]):
        line = image[lower:upper]

        pixel_values = processor(line, return_tensors="pt").pixel_values
        result = trocr_model.generate(pixel_values)

        text += processor.batch_decode(result, skip_special_tokens=True)[0]
        text += ' '

    return text


def add_text(document, text, text_type, is_handwritten):
    """
    Add text to docx file(document) depending on its type.
    Doesn't add figure and its annotation

    Args:
        document (docx.Document): document text will be added

        text (str): text to be added to document

        text_type (str): type of text from layoutparser model.
        Can be Text, Title or List.

        is_handwritten (bool): A flag indicating whether the text is handwritten.
        If `True`, the text will be italicized.

    Returns:
        None
    """

    if re.match('Figure [0-9]+\..*', text):
        return

    if is_handwritten:
        p = document.add_paragraph(' ')
        p.add_run(text).italic = True

    elif text_type == 'Text' or text_type == 'Figure':
        document.add_paragraph(text)

    elif text_type == 'Title':
        document.add_heading(text)

    elif text_type == 'List':
        # check whether layoutparser combine list elements
        for element in re.split('\* |\+ ', text):
            if element.strip() == '':
                continue
            document.add_paragraph(element, style='List Bullet')


def pdf_to_word(pdf_path):
    """
    Convert pdf file to docx file.

    Args:
        pdf_path (str): path to pdf file

    Returns:
        docx.Document: resulted docx file
    """

    doc = read_pdf(pdf_path)
    document = Document()

    main_progress_bar = tqdm(doc, desc='Document Processing')
    for page in main_progress_bar:
        prep_img = preprocess_image(page)

        text_blocks = model.detect(prep_img)

        text_blocks.sort(key=lambda b: b.coordinates[1], inplace=True)
        text_blocks = merge_blocks(text_blocks)

        for block in text_blocks:
            text, is_handwritten = get_text(block, prep_img)
            add_text(document, text, block.type, is_handwritten)

    return document


def digital_to_word(pdf_path):
    """
    Convert digital pdf file to docx file.

    Args:
        pdf_path (str): path to pdf file

    Returns:
        tuple:
            - docx.Document: A DOCX document object
            - bool: True if text extraction from the digital PDF was successful
    """

    try:
        doc = pymupdf.open(pdf_path)
    except:
        return None, False

    document = Document()

    is_readable = False

    for page in doc:
        blocks = page.get_text('blocks', sort=True)

        for block in blocks:
            text = block[4]
            document.add_paragraph(text)

            if text:
                is_readable = True

    return document, is_readable


def main(pdf_path, save=False):
    """
    Convert a PDF file to a DOCX file, using digital extraction
    if possible and fallback to OCR otherwise.

    Args:
        pdf_path (str): The file path to the PDF file to be converted.
        save (bool): Whether to save the resulting DOCX file.

    Returns:
        None
    """
    document, is_readable = digital_to_word(pdf_path)

    if not is_readable:
        document = pdf_to_word(pdf_path)

    if save:
        filename = pdf_path[:-4] + '.docx'
        document.save(filename)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Image preprocessing')

    parser.add_argument('--path', type=str,
                        default='results/scanned_mixed_handwritten_digital.pdf',
                        #default='results/scan_printed_text.pdf',
                        #default='results/digital.pdf',
                        help='path to pdf file')

    parser.add_argument('--save', type=bool, default=True,
                        help='whether to save docx file')

    args = parser.parse_args()

    main(args.path, save=args.save)
